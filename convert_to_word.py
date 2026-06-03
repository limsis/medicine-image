"""
将Markdown格式的技术报告转换为Word格式
"""
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def markdown_to_docx(input_file, output_file):
    """将Markdown文件转换为Word文档"""
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    # 读取Markdown文件
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 按行处理
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # 跳过空行
        if not line.strip():
            i += 1
            continue
        
        # 一级标题 (# )
        if line.startswith('# '):
            title = line[2:].strip()
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
        
        # 二级标题 (## )
        elif line.startswith('## '):
            title = line[3:].strip()
            doc.add_heading(title, level=2)
            i += 1
        
        # 三级标题 (### )
        elif line.startswith('### '):
            title = line[4:].strip()
            doc.add_heading(title, level=3)
            i += 1
        
        # 四级标题 (#### )
        elif line.startswith('#### '):
            title = line[5:].strip()
            doc.add_heading(title, level=4)
            i += 1
        
        # 无序列表 (- )
        elif line.startswith('- '):
            item = line[2:].strip()
            para = doc.add_paragraph(item, style='List Bullet')
            i += 1
        
        # 有序列表 (1. )
        elif re.match(r'^\d+\. ', line):
            item = re.sub(r'^\d+\. ', '', line).strip()
            para = doc.add_paragraph(item, style='List Number')
            i += 1
        
        # 代码块 (```)
        elif line.startswith('```'):
            code_block = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_block.append(lines[i])
                i += 1
            code_text = '\n'.join(code_block)
            para = doc.add_paragraph()
            run = para.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            para.paragraph_format.left_indent = Inches(0.5)
            i += 1
        
        # 行内代码 (`)
        elif '`' in line:
            # 处理行内代码
            parts = re.split(r'`([^`]+)`', line)
            para = doc.add_paragraph()
            for j, part in enumerate(parts):
                if j % 2 == 1:  # 代码部分
                    run = para.add_run(part)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 100, 0)
                else:  # 普通文本
                    para.add_run(part)
            i += 1
        
        # 链接和图片
        elif '](' in line:
            # 跳过链接和图片，保持文本
            text = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', r'\1', line)
            text = re.sub(r'\[([^\]]*)\]\([^\)]+\)', r'\1', text)
            if text.strip():
                doc.add_paragraph(text.strip())
            i += 1
        
        # 粗体 (**text**)
        elif '**' in line:
            parts = re.split(r'\*\*([^*]+)\*\*', line)
            para = doc.add_paragraph()
            for j, part in enumerate(parts):
                if j % 2 == 1:  # 粗体部分
                    run = para.add_run(part)
                    run.bold = True
                else:  # 普通文本
                    para.add_run(part)
            i += 1
        
        # 普通段落
        else:
            para = doc.add_paragraph(line.strip())
            i += 1
    
    # 保存文档
    doc.save(output_file)
    print(f"✅ 文档已保存到: {output_file}")

if __name__ == '__main__':
    input_file = '医学图像技术调研报告.md'
    output_file = '医学图像技术调研报告.docx'
    
    try:
        markdown_to_docx(input_file, output_file)
        print(f"🎉 转换成功！Word文档已生成：{output_file}")
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        print("\n提示：如果 python-docx 未安装，请运行：")
        print("pip install python-docx")
